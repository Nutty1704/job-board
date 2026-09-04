import json
import importlib.util
import sys
import unittest
from io import BytesIO
from pathlib import Path
from urllib.error import HTTPError


sys.path.insert(0, str(Path(__file__).resolve().parents[1]))
import job_resume  # noqa: E402


def profile_data():
    return {
        "version": "2026-08-25",
        "candidate_summary": "Early-career software engineer.",
        "filters": {"required_skills_any": ["python"], "allowed_locations": ["sydney"], "excluded_phrases": [], "max_required_experience_years": 3},
        "qualified_score_threshold": 80,
        "resume": {
            "identity": {"name": "Ada Example", "contact": "ada@example.test", "headline": "Software Engineer"},
            "summary": "Early-career software engineer.",
            "skill_catalog": {"Languages": ["Go", "Python"], "Cloud": ["AWS", "Kubernetes"]},
            "education": [{"institution": "Example University", "qualification": "Bachelor of Computer Science", "dates": "2020–2023"}],
            "additional_information": ["Australian permanent resident"],
            "experience": [
                {"id": "rokt", "kind": "technical", "employer": "Rokt", "title": "Software Engineer", "dates": "2024–Present", "source_bullets": [
                    {"id": "rokt-go", "text": "Shipped Go services.", "tags": ["go", "platform"]},
                    {"id": "rokt-k8s", "text": "Operated Kubernetes workloads.", "tags": ["kubernetes", "ownership"]},
                    {"id": "rokt-ci", "text": "Built CI pipelines.", "tags": ["ci/cd"]},
                    {"id": "rokt-api", "text": "Built APIs.", "tags": ["backend"]},
                    {"id": "rokt-data", "text": "Improved data reliability.", "tags": ["data"]},
                    {"id": "rokt-ops", "text": "Improved operations.", "tags": ["operations"]},
                ]},
                {"id": "retail", "kind": "transferable", "employer": "Shop", "title": "Team Member", "dates": "2020–2023", "source_bullets": [
                    {"id": "retail-team", "text": "Worked with customers.", "tags": ["communication"]},
                    {"id": "retail-own", "text": "Owned shift tasks.", "tags": ["ownership"]},
                ]},
            ],
            "projects": [{"id": "platform", "name": "Experiment Platform", "dates": "2025", "source_bullets": [
                {"id": "platform-fullstack", "text": "Built a full-stack experiment platform.", "tags": ["python", "react"]},
                {"id": "platform-oauth", "text": "Added OAuth SSO.", "tags": ["security"]},
                {"id": "platform-cicd", "text": "Automated platform deployment.", "tags": ["ci/cd"]},
            ]}],
        },
    }


def high_match():
    return {"source": "adzuna", "source_job_id": "123", "profile_s3_version": "profile-v1", "job_event": {"job": {"title": "Platform Engineer", "description": "Go and Kubernetes."}}}


class ProfileValidationTests(unittest.TestCase):
    def test_parse_profile_requires_tagged_resume_source_bullets(self):
        value = profile_data()
        del value["resume"]["experience"][0]["source_bullets"][0]["tags"]

        with self.assertRaisesRegex(ValueError, "tags"):
            job_resume.parse_resume_profile(json.dumps(value).encode())

    def test_validate_selection_rejects_invalid_source_bullet_sets(self):
        profile = job_resume.parse_resume_profile(json.dumps(profile_data()).encode())
        selection = {"summary": "Early-career software engineer.", "skill_groups": [{"group": "Languages", "skills": ["Go"]}], "experience": [{"id": "rokt", "source_bullet_ids": ["rokt-go", "rokt-k8s", "rokt-ci", "rokt-api", "rokt-data", "rokt-ops"]}, {"id": "retail", "source_bullet_ids": ["retail-team"]}], "projects": []}

        with self.assertRaisesRegex(ValueError, "technical experience"):
            job_resume.validate_selection(selection, profile)

        selection["experience"][0]["source_bullet_ids"] = ["rokt-go", "rokt-k8s", "rokt-ci", "rokt-api"]

        with self.assertRaisesRegex(ValueError, "transferable experience"):
            job_resume.validate_selection(selection, profile)

        selection["experience"][1]["source_bullet_ids"] = ["retail-team", "retail-own"]
        selection["experience"][0]["source_bullet_ids"] = ["rokt-go", "rokt-go", "rokt-ci", "rokt-api"]
        with self.assertRaisesRegex(ValueError, "duplicate source bullet"):
            job_resume.validate_selection(selection, profile)

        selection["experience"][0]["source_bullet_ids"] = ["rokt-go", "rokt-k8s", "rokt-ci", "unknown"]
        with self.assertRaisesRegex(ValueError, "unknown source bullet"):
            job_resume.validate_selection(selection, profile)

    def test_selection_schema_uses_openai_compatible_project_identifier_enums(self):
        profile = job_resume.parse_resume_profile(json.dumps(profile_data()).encode())

        schema = job_resume._selection_schema(profile)
        experience = schema["properties"]["experience"]
        rokt = experience["properties"]["rokt"]
        retail = experience["properties"]["retail"]

        self.assertEqual(experience["required"], ["rokt", "retail"])
        self.assertFalse(experience["additionalProperties"])
        self.assertEqual(rokt["minItems"], 4)
        self.assertEqual(rokt["maxItems"], 5)
        self.assertEqual(rokt["items"]["enum"], ["rokt-go", "rokt-k8s", "rokt-ci", "rokt-api", "rokt-data", "rokt-ops"])
        self.assertEqual(retail["minItems"], 2)
        self.assertEqual(retail["maxItems"], 2)
        projects = schema["properties"]["projects"]
        self.assertEqual(projects["maxItems"], 3)
        project = projects["items"]["anyOf"][0]
        self.assertEqual(project["properties"]["id"]["enum"], ["platform"])
        self.assertNotIn("const", project["properties"]["id"])
        bullets = project["properties"]["source_bullet_ids"]
        self.assertEqual(bullets["minItems"], 3)
        self.assertEqual(bullets["maxItems"], 3)
        self.assertEqual(bullets["items"]["enum"], ["platform-fullstack", "platform-oauth", "platform-cicd"])

    def test_selection_schema_limits_skill_groups_to_the_profile_catalog(self):
        profile = job_resume.parse_resume_profile(json.dumps(profile_data()).encode())

        groups = job_resume._selection_schema(profile)["properties"]["skill_groups"]["items"]

        self.assertIn("anyOf", groups)
        by_name = {item["properties"]["group"]["enum"][0]: item for item in groups["anyOf"]}
        self.assertEqual(set(by_name), {"Languages", "Cloud"})
        self.assertEqual(by_name["Languages"]["properties"]["skills"]["items"]["enum"], ["Go", "Python"])
        self.assertEqual(by_name["Cloud"]["properties"]["skills"]["items"]["enum"], ["AWS", "Kubernetes"])

class FakeS3:
    def __init__(self, existing_keys=None):
        self.puts = []
        self.existing_keys = set(existing_keys or [])

    def get_object(self, **kwargs):
        if kwargs["Key"] == "matching/current.json":
            return {"Body": Body(json.dumps(profile_data()).encode()), "VersionId": "profile-v1"}
        return {"Body": Body(b"template"), "VersionId": "template-v2"}

    def put_object(self, **kwargs):
        self.puts.append(kwargs)
        self.existing_keys.add(kwargs["Key"])
        return {"VersionId": "artifact-v3"}

    def head_object(self, **kwargs):
        if kwargs["Key"] not in self.existing_keys:
            raise FileNotFoundError(kwargs["Key"])
        return {"VersionId": "artifact-v3"}


class Body:
    def __init__(self, value): self.value = value
    def read(self): return self.value


class FakeParameters:
    def get_parameter(self, **kwargs): return {"Parameter": {"Value": '{"api_key":"test-key"}'}}


class ConsumerTests(unittest.TestCase):
    def selection(self):
        return {"summary": "Early-career software engineer.", "skill_groups": [{"group": "Languages", "skills": ["Go"]}], "experience": [{"id": "rokt", "source_bullet_ids": ["rokt-go", "rokt-k8s", "rokt-ci", "rokt-api"]}, {"id": "retail", "source_bullet_ids": ["retail-team", "retail-own"]}], "projects": [{"id": "platform", "source_bullet_ids": ["platform-fullstack", "platform-oauth", "platform-cicd"]}]}

    def consumer(self, s3=None, request=None):
        return job_resume.ResumeConsumer(job_resume.Config.from_environment({}), s3 or FakeS3(), FakeParameters(), request or (lambda *_: {"id": "resp-1", "output_text": json.dumps(self.selection())}), lambda template, context: b"docx")

    def test_consumer_stores_rendered_artifact_at_job_scoped_s3_key(self):
        requests = []
        consumer = self.consumer(request=lambda request: requests.append(request) or {"id": "resp-1", "output_text": json.dumps(self.selection())})
        result = consumer.process(high_match())

        self.assertEqual(result, "completed")
        upload = consumer.s3.puts[0]
        self.assertEqual(upload["Key"], "resumes/123.docx")
        self.assertEqual(upload["Metadata"]["profile-s3-version"], "profile-v1")
        self.assertEqual(upload["Metadata"]["source"], "adzuna")
        self.assertEqual(upload["Metadata"]["source-job-id"], "123")
        self.assertEqual(upload["Tagging"], "expires-after-days=21")
        self.assertTrue(requests[0]["text"]["format"]["strict"])

    def test_existing_job_resume_skips_model_call(self):
        calls = []
        consumer = self.consumer(FakeS3({"resumes/123.docx"}), lambda *_: calls.append(True))

        self.assertEqual(consumer.process(high_match()), "duplicate")
        self.assertEqual(calls, [])

    def test_manual_resume_uses_a_namespaced_artifact_key(self):
        message = {**high_match(), "source": "manual", "source_job_id": "123e4567-e89b-12d3-a456-426614174000"}
        consumer = self.consumer()

        self.assertEqual(consumer.process(message), "completed")
        self.assertEqual(consumer.s3.puts[0]["Key"], "resumes/manual/123e4567-e89b-12d3-a456-426614174000.docx")

    def test_openai_request_error_includes_response_detail(self):
        original_urlopen = job_resume.urlopen
        job_resume.urlopen = lambda *_args, **_kwargs: (_ for _ in ()).throw(HTTPError("https://api.openai.com/v1/responses", 400, "Bad Request", {}, BytesIO(b'{"error":{"message":"invalid_json_schema"}}')))
        try:
            with self.assertRaisesRegex(RuntimeError, "invalid_json_schema"):
                job_resume.openai_response("secret", "gpt-5.6-luna", {"job": {}})
        finally:
            job_resume.urlopen = original_urlopen

    @unittest.skipUnless(importlib.util.find_spec("docxtpl"), "requires DOCX rendering dependencies")
    def test_render_docx_preserves_ampersands_in_dynamic_text(self):
        from docx import Document

        template = BytesIO()
        document = Document()
        document.add_paragraph("{{ group }}")
        document.save(template)

        rendered = Document(BytesIO(job_resume.render_docx(template.getvalue(), {"group": "Cloud & Platform"})))

        self.assertEqual(rendered.paragraphs[0].text, "Cloud & Platform")

    def test_batch_retries_only_active_or_failed_work(self):
        response = job_resume.process_sqs_batch({"Records": [{"messageId": "active", "body": json.dumps(high_match())}]}, type("C", (), {"process": lambda self, _: "retry"})())
        self.assertEqual(response, {"batchItemFailures": [{"itemIdentifier": "active"}]})
